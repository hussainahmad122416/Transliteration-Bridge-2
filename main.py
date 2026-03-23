from fastapi import FastAPI, UploadFile, File, HTTPException, Form, Request
from fastapi.staticfiles import StaticFiles
from fastapi.responses import HTMLResponse, FileResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded
import time
import io
import os
import uvicorn
from docx import Document
import re
import PyPDF2
from functools import lru_cache

# Custom Modules
from modules.preprocessing import PreprocessingAgent
from modules.transliteration import TransliterationEngine
from modules.post_processing import PostProcessingAgent

# DB & Routing imports
from backend.database import engine, Base
from backend.auth import auth_router
from backend.history import history_router

# Initialize SQLite Database tables
Base.metadata.create_all(bind=engine)

limiter = Limiter(key_func=get_remote_address)
app = FastAPI(title="Transliteration Bridge API v2")
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

from gtts import gTTS

# Include Auth and History APIS
app.include_router(auth_router)
app.include_router(history_router)

# Initialize Pipeline Modules
preprocessor = PreprocessingAgent()
translator = TransliterationEngine()
post_processor = PostProcessingAgent(dictionary_path="data/dictionary.json")

# Ensure static directory exists
os.makedirs("static", exist_ok=True)
app.mount("/static", StaticFiles(directory="static"), name="static")

@app.get("/")
async def read_index():
    return FileResponse("static/index.html")

class TextRequest(BaseModel):
    text: str = Field(..., max_length=5000, description="The raw text to transliterate, max 5000 chars to prevent DoS.")
    direction: str = "shahmukhi_to_gurmukhi" # default

class TTSRequest(BaseModel):
    text: str = Field(..., max_length=1500)
    lang: str = "ur"

@lru_cache(maxsize=1000)
def run_transliteration_pipeline(raw_text: str, direction: str) -> dict:
    start_time = time.time()
    
    # 1. Preprocessing (Clean and tokenize)
    prep_res = preprocessor.process(raw_text)
    tokens = prep_res['tokens']
    
    # 2. Transliteration (Rule-based bidirectional)
    transliterated_tokens = translator.process(tokens, direction=direction)
    
    # 3. Post-Processing (Dictionary override)
    corrected_tokens = post_processor.apply_corrections(tokens, transliterated_tokens, direction=direction)
    
    # 4. Final Output Construction
    import string
    
    # Simple reconstruction logic
    result_text = ""
    rich_mapping = []
    
    for token, original in zip(corrected_tokens, tokens):
        # Match words considering Shahmukhi and Gurmukhi specific unicode blocks
        is_word = bool(re.match(r'^[\w\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB50-\uFDFF\uFE70-\uFEFF\u0A00-\u0A7F\u200C]+$', original, re.UNICODE))
        
        if is_word:
            if result_text and not result_text.endswith(' '):
                result_text += ' '
                rich_mapping.append({"is_word": False, "text": " "})
                
        result_text += token
        
        if is_word:
            rich_mapping.append({
                "is_word": True,
                "input": original,
                "output": token
            })
        else:
            rich_mapping.append({
                "is_word": False,
                "text": original
            })
        
    result_text = result_text.strip()
    
    # Recalculate stats for Gurmukhi Output
    word_count = len([t for t in result_text.split() if any(c.isalnum() for c in t)])
    char_count = len(result_text)
    process_time_ms = int((time.time() - start_time) * 1000)
    
    return {
        "input_text": raw_text,
        "output_text": result_text,
        "direction": direction,
        "word_count": word_count,
        "char_count": char_count,
        "process_time_ms": process_time_ms,
        "is_source_script_correct": prep_res['is_shahmukhi'] if direction == "shahmukhi_to_gurmukhi" else True,
        "rich_mapping": rich_mapping
    }

import re
@app.post("/api/transliterate")
@limiter.limit("20/minute")
async def transliterate_text(payload: TextRequest, request: Request):
    try:
        if not payload.text.strip():
            return {"output_text": "", "word_count": 0, "char_count": 0, "process_time_ms": 0}
            
        result = run_transliteration_pipeline(payload.text, payload.direction)
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/transliterate/file")
@limiter.limit("10/minute")
async def transliterate_file(request: Request, direction: str = Form("shahmukhi_to_gurmukhi"), file: UploadFile = File(...)):
    try:
        contents = await file.read()
        text = ""
        
        if file.filename.endswith(".txt"):
            text = contents.decode("utf-8")
        elif file.filename.endswith(".docx"):
            doc = Document(io.BytesIO(contents))
            text = "\n".join([para.text for para in doc.paragraphs])
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type")
            
        result = run_transliteration_pipeline(text, direction)
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/parse-file")
@limiter.limit("15/minute")
async def parse_file(request: Request, file: UploadFile = File(...)):
    try:
        contents = await file.read()
        text = ""
        filename = file.filename.lower()
        
        if filename.endswith(".txt"):
            text = contents.decode("utf-8")
        elif filename.endswith(".docx"):
            doc = Document(io.BytesIO(contents))
            text = "\n".join([para.text for para in doc.paragraphs])
        elif filename.endswith(".pdf"):
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(contents))
            for page in pdf_reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
        else:
            raise HTTPException(status_code=400, detail="Unsupported file format")
            
        # Basic Language Detection based on Unicode blocks
        # Shahmukhi: Arabic blocks (0x0600 - 0x06FF)
        # Gurmukhi: Gurmukhi block (0x0A00 - 0x0A7F)
        shahmukhi_count = len(re.findall(r'[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF]', text))
        gurmukhi_count = len(re.findall(r'[\u0A00-\u0A7F]', text))
        
        detected_lang = "unknown"
        if shahmukhi_count > gurmukhi_count and shahmukhi_count > 10:
            detected_lang = "shahmukhi"
        elif gurmukhi_count > shahmukhi_count and gurmukhi_count > 10:
            detected_lang = "gurmukhi"
            
        return {
            "text": text.strip(),
            "detected_language": detected_lang,
            "filename": file.filename
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/tts")
@limiter.limit("30/minute")
async def text_to_speech(request: Request, payload: TTSRequest):
    try:
        if not payload.text.strip():
            raise HTTPException(status_code=400, detail="Empty text")
        # Generate MP3 using Google TTS
        tts = gTTS(text=payload.text, lang=payload.lang)
        fp = io.BytesIO()
        tts.write_to_fp(fp)
        fp.seek(0)
        return StreamingResponse(fp, media_type="audio/mpeg")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/export-docx")
@limiter.limit("15/minute")
async def export_docx(request: Request, payload: TextRequest):
    try:
        if not payload.text.strip():
            raise HTTPException(status_code=400, detail="Empty text")
        
        doc = Document()
        doc.add_paragraph(payload.text)
        
        fp = io.BytesIO()
        doc.save(fp)
        fp.seek(0)
        
        headers = {
            'Content-Disposition': f'attachment; filename="transliteration_{payload.direction}.docx"'
        }
        return StreamingResponse(fp, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers=headers)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    uvicorn.run("main:app", host="127.0.0.1", port=8000, reload=True)
